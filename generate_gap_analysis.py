"""
Executive Dashboard Gap Analysis Document Generator
Generates a comprehensive Word document analyzing the current state and gaps
for transforming the dashboard into a Board-level Executive Decision Dashboard
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_formatting(doc, text, bold=False, italic=False, font_size=11):
    """Add a formatted paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return para

def add_bullet_point(doc, text, level=0):
    """Add a bullet point with optional indentation"""
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.5 * level)
    return para

def create_gap_analysis_document():
    """Generate the complete gap analysis document"""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ========================================
    # TITLE PAGE
    # ========================================
    title = doc.add_heading('Executive Dashboard Gap Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Transforming Contract Oversight into a Board-Level Executive Decision Dashboard')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(68, 84, 106)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    client_info = doc.add_paragraph()
    client_run = client_info.add_run('Marion County School District\nCapital Projects & Infrastructure Oversight')
    client_run.font.size = Pt(12)
    client_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(10)
    date_run.font.italic = True
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========================================
    # 1. EXECUTIVE SUMMARY
    # ========================================
    add_heading_with_style(doc, '1. Executive Summary', level=1)

    doc.add_paragraph(
        'This gap analysis evaluates the current Marion County Contract Oversight System and identifies '
        'critical transformations needed to create a high-impact Executive Decision Dashboard for School '
        'Board members, County Commissioners, and senior leadership overseeing capital projects and infrastructure investments.'
    )

    add_heading_with_style(doc, 'What Is Happening Now', level=2)
    doc.add_paragraph(
        'The existing system is a comprehensive, operational-level contract management platform containing '
        '38+ features, extensive data tracking, and detailed analytics. While technically robust, the current '
        'dashboard presents significant information density that overwhelms executive-level users who need '
        'rapid, strategic insights during board meetings.'
    )

    add_heading_with_style(doc, 'What Matters Most', level=2)

    add_bullet_point(doc, 'Board members need to understand portfolio health in under 5 minutes')
    add_bullet_point(doc, 'Critical decision points must be immediately visible (project risks, budget overruns, timeline delays)')
    add_bullet_point(doc, 'Land costs must be separated from construction costs for accurate capital project analysis')
    add_bullet_point(doc, 'Real data is required—current sample data undermines credibility and decision-making')
    add_bullet_point(doc, 'Capital projects require specialized risk assessment and project management metrics')
    add_bullet_point(doc, 'Historical vendor performance over 20+ years is needed for strategic vendor selection')

    add_heading_with_style(doc, 'Why It Matters', level=2)
    doc.add_paragraph(
        'School board executives and county commissioners make multi-million dollar funding decisions based '
        'on the information presented. Unclear or overwhelming dashboards lead to:'
    )

    add_bullet_point(doc, 'Delayed decision-making due to information overload')
    add_bullet_point(doc, 'Missed warning signs of project failures or budget overruns')
    add_bullet_point(doc, 'Inability to compare projects meaningfully (e.g., mixing land and construction costs)')
    add_bullet_point(doc, 'Reduced trust in the system when sample data is presented as real')
    add_bullet_point(doc, 'Lost opportunities for strategic vendor negotiations and long-term planning')

    add_heading_with_style(doc, 'Immediate Executive-Level Takeaways', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Current State: ')
    run.bold = True
    para.add_run('The dashboard is an operational management tool, not an executive briefing tool.')

    para = doc.add_paragraph()
    run = para.add_run('Primary Gap: ')
    run.bold = True
    para.add_run(
        'Information hierarchy is inverted—operational details overshadow strategic insights. '
        'Executives cannot quickly answer: "What needs my attention?" and "What decision must I make?"'
    )

    para = doc.add_paragraph()
    run = para.add_run('Critical Risk: ')
    run.bold = True
    para.add_run(
        'Without transformation, board members will continue to request manual briefings, '
        'reducing dashboard adoption and undermining the system\'s ROI.'
    )

    para = doc.add_paragraph()
    run = para.add_run('Opportunity: ')
    run.bold = True
    para.add_run(
        'With focused refinement, this system can become the authoritative source for all board-level '
        'capital project decisions, setting a best-practice standard for government transparency.'
    )

    doc.add_page_break()

    # ========================================
    # 2. CURRENT STATE ASSESSMENT
    # ========================================
    add_heading_with_style(doc, '2. Current State Assessment', level=1)

    add_heading_with_style(doc, 'What the Dashboard Currently Shows', level=2)

    doc.add_paragraph(
        'The Marion County Contract Oversight System is a fully-functional, production-ready platform with '
        'comprehensive contract lifecycle management capabilities:'
    )

    add_heading_with_style(doc, 'Dashboard Components:', level=3)

    add_bullet_point(doc, 'Two primary dashboards: Standard Dashboard (operational) and Executive Dashboard (board-level)')
    add_bullet_point(doc, 'Overall portfolio health score (0-100 scale with color-coded status)')
    add_bullet_point(doc, 'Financial metrics: Total contract value, amount paid, cost overruns, budget variance')
    add_bullet_point(doc, 'Performance tracking: On-time delivery rate, vendor scores, milestone completion')
    add_bullet_point(doc, 'Risk indicators: At-risk contracts, critical issues, expiring contracts')
    add_bullet_point(doc, 'Vendor analytics: Diversity metrics, local vendor participation, performance ratings')
    add_bullet_point(doc, 'Compliance tracking: Insurance verification, board approvals, competitive bidding')
    add_bullet_point(doc, 'Benchmarking against Coupa 2025 industry standards (20 KPIs across 8 categories)')
    add_bullet_point(doc, 'Peer county comparison using Florida EDR data')

    add_heading_with_style(doc, 'Types of Metrics, Visuals, and Data Density', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Metrics Tracked: ')
    run.bold = True
    para.add_run('The system tracks over 50 distinct metrics across financial, operational, compliance, and performance categories.')

    para = doc.add_paragraph()
    run = para.add_run('Visualization Types: ')
    run.bold = True
    para.add_run('Pie charts, bar charts, line charts, progress bars, radial/donut charts—with drill-down and export capabilities.')

    para = doc.add_paragraph()
    run = para.add_run('Data Density: ')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    para.add_run(
        'HIGH. The executive dashboard displays 6 quick stats, 4 health score metrics, 3 priority sections, '
        '2 major charts, 6 KPIs, school district summary, peer comparison, and top contracts table—all on one page. '
        'This exceeds the cognitive load threshold for executive consumption.'
    )

    add_heading_with_style(doc, 'Intended vs Actual Usability for Executives', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Intended Use: ')
    run.bold = True
    para.add_run('Provide board-level overview for strategic decision-making during meetings.')

    para = doc.add_paragraph()
    run = para.add_run('Actual Use: ')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    para.add_run(
        'The current design requires 15-20 minutes to fully digest. Board members in a meeting setting '
        'will struggle to identify priorities, understand causation, or make informed decisions within typical '
        'agenda timeframes (5-10 minutes per topic).'
    )

    doc.add_paragraph(
        'Key usability issues:'
    )
    add_bullet_point(doc, 'No clear visual hierarchy—all information appears equally important')
    add_bullet_point(doc, 'Lacks narrative flow—data is presented without context or interpretation')
    add_bullet_point(doc, 'Missing "so what?" factor—charts show data but don\'t tell executives what action to take')
    add_bullet_point(doc, 'No predictive indicators—all metrics are backward-looking (what happened, not what will happen)')
    add_bullet_point(doc, 'Limited decision support—doesn\'t answer "Should we approve this?" or "What is the risk if we proceed?"')

    add_heading_with_style(doc, 'Strengths That Should Be Retained', level=2)

    doc.add_paragraph('The current system has significant strengths that form a solid foundation:')

    add_bullet_point(doc, 'Sophisticated health scoring algorithm (4-dimensional: cost, schedule, performance, compliance)')
    add_bullet_point(doc, 'Comprehensive data model with proper audit trails and change tracking')
    add_bullet_point(doc, 'Industry benchmarking integration (Coupa 2025 standards)')
    add_bullet_point(doc, 'Export capabilities (PDF, Excel, print-friendly formats)')
    add_bullet_point(doc, 'Real-time alert generation with severity-based prioritization')
    add_bullet_point(doc, 'Vendor performance tracking and diversity metrics')
    add_bullet_point(doc, 'Peer county comparison framework')
    add_bullet_point(doc, 'Dark mode support and responsive design')
    add_bullet_point(doc, 'Approval workflow and governance features')
    add_bullet_point(doc, 'Strong technical architecture (Flask, SQLite, Plotly, Tailwind CSS)')

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Critical Insight: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    para.add_run(
        'The underlying data infrastructure and analytics engine are executive-grade. The gap is purely '
        'in presentation layer and information architecture—not in capability or data quality.'
    )

    doc.add_page_break()

    # ========================================
    # 3. GAP ANALYSIS
    # ========================================
    add_heading_with_style(doc, '3. Gap Analysis', level=1)

    doc.add_paragraph(
        'This section identifies critical gaps across five dimensions, with each gap assessed for '
        'board-level impact and implementation risk.'
    )

    doc.add_paragraph()

    # Gap 1: Information Hierarchy
    add_heading_with_style(doc, 'Gap 1: Information Hierarchy', level=2)

    para = doc.add_paragraph()
    run = para.add_run('What Executives Need:')
    run.bold = True

    add_bullet_point(doc, 'Immediate answer to: "What requires my decision today?"')
    add_bullet_point(doc, 'Top 3-5 projects/issues ranked by strategic importance')
    add_bullet_point(doc, 'Red/yellow/green status at-a-glance (no need to interpret numbers)')
    add_bullet_point(doc, 'Exception-based reporting (only show what\'s outside acceptable parameters)')
    add_bullet_point(doc, 'One-number summaries with drill-down capability for details')

    para = doc.add_paragraph()
    run = para.add_run('What Is Currently Shown:')
    run.bold = True

    add_bullet_point(doc, 'All contracts weighted equally regardless of strategic importance')
    add_bullet_point(doc, 'Numerical metrics requiring interpretation (e.g., "Health Score: 67"—is this good or bad?)')
    add_bullet_point(doc, 'Comprehensive data with no prioritization or filtering')
    add_bullet_point(doc, 'Flat information architecture—operational and strategic data mixed together')

    para = doc.add_paragraph()
    run = para.add_run('Why It Matters at Board Level:')
    run.bold = True

    doc.add_paragraph(
        'Board members allocate 5-10 minutes per agenda item. If they cannot identify the #1 issue requiring '
        'their attention within 30 seconds, the dashboard has failed its primary purpose. Currently, executives '
        'must mentally process multiple data points to derive priority—this cognitive burden leads to decision '
        'fatigue and reduced dashboard adoption.'
    )

    para = doc.add_paragraph()
    run = para.add_run('Risk of Not Addressing:')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    add_bullet_point(doc, 'Board members will revert to requesting manual PowerPoint briefings')
    add_bullet_point(doc, 'Critical project failures may be missed in the noise of comprehensive data')
    add_bullet_point(doc, 'Decision paralysis—too much information leads to delayed or avoided decisions')
    add_bullet_point(doc, 'System ROI undermined—investment in dashboard development not realized')

    doc.add_paragraph()

    # Gap 2: Visual Design & Layout
    add_heading_with_style(doc, 'Gap 2: Visual Design & Layout (Clarity, Cognitive Load, Narrative Flow)', level=2)

    para = doc.add_paragraph()
    run = para.add_run('What Is Missing or Misaligned:')
    run.bold = True

    add_bullet_point(doc, 'No guided narrative—dashboard presents data without telling a story')
    add_bullet_point(doc, 'Excessive use of charts—multiple visualizations compete for attention')
    add_bullet_point(doc, 'No white space strategy—dense layouts increase cognitive load')
    add_bullet_point(doc, 'Color coding inconsistent—red/yellow/green used but not as primary navigation')
    add_bullet_point(doc, 'Missing summary cards—no "executive snapshot" summarizing overall status')
    add_bullet_point(doc, 'No progressive disclosure—all detail visible at once rather than summary → drill-down')
    add_bullet_point(doc, 'Lack of annotations—charts show trends but don\'t explain causation')

    para = doc.add_paragraph()
    run = para.add_run('Why It Matters at Board Level:')
    run.bold = True

    doc.add_paragraph(
        'Executives process information differently than operational managers. They need high-level patterns, '
        'not granular data. The current dashboard violates cognitive load principles by presenting all information '
        'simultaneously. Best practices for executive dashboards recommend:'
    )

    add_bullet_point(doc, '3-5 primary metrics maximum on first screen')
    add_bullet_point(doc, 'Story-driven layout (beginning → middle → end)')
    add_bullet_point(doc, 'Annotations explaining "why" metrics changed')
    add_bullet_point(doc, 'Clear visual hierarchy (most important = largest/top-left)')

    para = doc.add_paragraph()
    run = para.add_run('Risk of Not Addressing:')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    add_bullet_point(doc, 'Information overload leads to executives ignoring the dashboard entirely')
    add_bullet_point(doc, 'Misinterpretation of data due to lack of context')
    add_bullet_point(doc, 'Inability to communicate insights to constituents or media')
    add_bullet_point(doc, 'Dashboard credibility damaged—perceived as "too technical" or "not user-friendly"')

    doc.add_paragraph()

    # Gap 3: Executive KPIs
    add_heading_with_style(doc, 'Gap 3: Executive KPIs (Missing Strategic Indicators)', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Critical Missing KPIs:')
    run.bold = True

    add_bullet_point(doc, 'Land Cost vs Construction Cost Separation—currently blended, making project comparisons impossible')
    add_bullet_point(doc, 'Project ROI and Long-Term Value—no analysis of benefits delivered per dollar invested')
    add_bullet_point(doc, 'Expected Completion Date Based on Current Burn Rate—no forward-looking projections')
    add_bullet_point(doc, 'Budget Forecast to Completion—no estimate of final project costs based on current trajectory')
    add_bullet_point(doc, 'Vendor Concentration Risk—no visibility into over-reliance on single vendors')
    add_bullet_point(doc, 'Historical Vendor Performance (20+ years)—no long-term vendor track record')
    add_bullet_point(doc, 'Capital Project Pipeline—no view of upcoming projects by fiscal year')
    add_bullet_point(doc, 'Risk Likelihood Scoring—current system shows current risk, not probability of future failure')
    add_bullet_point(doc, 'Community Impact Metrics—no measurement of projects by district/constituent impact')
    add_bullet_point(doc, 'Inflation-Adjusted Historical Costs—no dollar normalization for accurate comparisons')

    para = doc.add_paragraph()
    run = para.add_run('Why It Matters at Board Level:')
    run.bold = True

    doc.add_paragraph(
        'Board decisions are inherently strategic and forward-looking. Current KPIs are operational and backward-looking. '
        'Executives need to answer questions like:'
    )

    add_bullet_point(doc, '"Can we afford to start the new school construction project next fiscal year?"')
    add_bullet_point(doc, '"Is this vendor reliable for a 20-year facilities management contract?"')
    add_bullet_point(doc, '"Which district has received the most capital investment over the past decade?"')
    add_bullet_point(doc, '"What is the risk this project will fail, and what would be the financial impact?"')

    doc.add_paragraph('None of these questions can currently be answered from the dashboard.')

    para = doc.add_paragraph()
    run = para.add_run('Risk of Not Addressing:')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    add_bullet_point(doc, 'Poor strategic decisions due to incomplete information')
    add_bullet_point(doc, 'Inability to justify budget requests to state agencies or taxpayers')
    add_bullet_point(doc, 'Vendor selection based on incomplete performance history')
    add_bullet_point(doc, 'Inequitable resource distribution across school districts')
    add_bullet_point(doc, 'Project failures that could have been predicted with proper risk scoring')

    doc.add_paragraph()

    # Gap 4: Decision Enablement
    add_heading_with_style(doc, 'Gap 4: Decision Enablement (What Decisions Cannot Currently Be Made)', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Decisions Requiring Dashboard Support (Currently Not Enabled):')
    run.bold = True

    doc.add_paragraph('1. Funding Approval Decisions')
    add_bullet_point(doc, '"Should we approve the $50M new elementary school project?"', level=1)
    add_bullet_point(doc, 'Requires: Land cost breakdown, comparable project benchmarks, vendor track record, risk assessment', level=1)
    add_bullet_point(doc, 'Currently missing: All of the above', level=1)

    doc.add_paragraph('2. Vendor Selection and Negotiation')
    add_bullet_point(doc, '"Which vendor should we award the 10-year HVAC maintenance contract to?"', level=1)
    add_bullet_point(doc, 'Requires: 20-year performance history, price trends, reliability metrics, capacity analysis', level=1)
    add_bullet_point(doc, 'Currently missing: Historical performance beyond current contracts', level=1)

    doc.add_paragraph('3. Project Continuation or Termination')
    add_bullet_point(doc, '"Should we halt this project or authorize additional funding to complete it?"', level=1)
    add_bullet_point(doc, 'Requires: Cost to complete forecast, risk of failure analysis, alternative options comparison', level=1)
    add_bullet_point(doc, 'Currently missing: Predictive forecasting and scenario modeling', level=1)

    doc.add_paragraph('4. Multi-Year Capital Planning')
    add_bullet_point(doc, '"What is our capital project capacity for the next 5 fiscal years?"', level=1)
    add_bullet_point(doc, 'Requires: Pipeline view, historical spend patterns, inflation adjustments, revenue projections', level=1)
    add_bullet_point(doc, 'Currently missing: Pipeline view and inflation-adjusted historical data', level=1)

    doc.add_paragraph('5. Constituent Equity and Transparency')
    add_bullet_point(doc, '"Are we distributing capital investments fairly across all school districts?"', level=1)
    add_bullet_point(doc, 'Requires: Per-student spending by district, geographic visualization, historical equity analysis', level=1)
    add_bullet_point(doc, 'Currently missing: Granular district-level analysis and geographic views', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Why It Matters at Board Level:')
    run.bold = True

    doc.add_paragraph(
        'A dashboard that cannot support actual board decisions is merely a reporting tool, not a decision support system. '
        'The value proposition of this investment is predicated on enabling faster, better-informed decisions. Without '
        'decision enablement features, the dashboard becomes a "nice to have" rather than a "must have."'
    )

    para = doc.add_paragraph()
    run = para.add_run('Risk of Not Addressing:')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    add_bullet_point(doc, 'Board continues using external consultants for decision support (ongoing cost, slower decisions)')
    add_bullet_point(doc, 'Decisions made with incomplete data, leading to suboptimal outcomes')
    add_bullet_point(doc, 'Public perception of lack of transparency ("Why can\'t we see how taxpayer money is spent?")')
    add_bullet_point(doc, 'Legal/compliance risk from inequitable resource distribution')

    doc.add_paragraph()

    # Gap 5: Narrative & Storytelling
    add_heading_with_style(doc, 'Gap 5: Narrative & Storytelling (Lack of Context, Trends, or Insights)', level=2)

    para = doc.add_paragraph()
    run = para.add_run('What Is Missing:')
    run.bold = True

    add_bullet_point(doc, 'No executive summary text—dashboard is 100% visualizations and data, 0% narrative')
    add_bullet_point(doc, 'No trend analysis—shows current state but not "How did we get here?" or "Where are we headed?"')
    add_bullet_point(doc, 'No automated insights—system doesn\'t explain why metrics changed (e.g., "Budget variance increased due to 3 change orders in Q2")')
    add_bullet_point(doc, 'No comparative context—no "This is X% better/worse than last year" or "We rank #3 among peer counties"')
    add_bullet_point(doc, 'No success stories highlighted—positive outcomes buried in comprehensive data')
    add_bullet_point(doc, 'No recommendations—system identifies problems but doesn\'t suggest solutions')
    add_bullet_point(doc, 'No scenario planning—can\'t model "What if we delay this project 6 months?" or "What if costs increase 15%?"')

    para = doc.add_paragraph()
    run = para.add_run('Why It Matters at Board Level:')
    run.bold = True

    doc.add_paragraph(
        'Executives are storytellers. They must explain decisions to constituents, media, and oversight agencies. '
        'A dashboard that presents data without narrative context forces executives to manually construct the story, '
        'which is time-consuming and risks misinterpretation. Best-in-class executive dashboards provide:'
    )

    add_bullet_point(doc, 'Auto-generated executive summaries ("This month, 3 projects exceeded budget due to material cost increases")')
    add_bullet_point(doc, 'Annotated charts explaining inflection points ("Why did spending spike in March?")')
    add_bullet_point(doc, 'Comparative statements ("We\'re performing 12% better than peer counties on on-time delivery")')
    add_bullet_point(doc, 'Recommendations based on data patterns ("Consider renegotiating with Vendor X—performance declining for 3 quarters")')

    para = doc.add_paragraph()
    run = para.add_run('Risk of Not Addressing:')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    add_bullet_point(doc, 'Executives misinterpret data, leading to incorrect public statements')
    add_bullet_point(doc, 'Inability to effectively communicate with taxpayers and media')
    add_bullet_point(doc, 'Low dashboard engagement—without narrative, executives perceive it as "just numbers"')
    add_bullet_point(doc, 'Missed opportunities to celebrate successes and build public trust')

    doc.add_page_break()

    # ========================================
    # 4. IDEAL EXECUTIVE DASHBOARD VISION
    # ========================================
    add_heading_with_style(doc, '4. Ideal Executive Dashboard Vision', level=1)

    add_heading_with_style(doc, 'What an Optimal Board-Level Dashboard Should Look Like', level=2)

    doc.add_paragraph(
        'An executive-grade decision dashboard for school board capital projects should embody these principles:'
    )

    add_heading_with_style(doc, 'Principle 1: Clarity Over Comprehensiveness', level=3)
    doc.add_paragraph(
        'Show the vital few, not the trivial many. The dashboard should answer 3 questions:'
    )
    add_bullet_point(doc, 'What is the overall status? (One number, color-coded: Green/Yellow/Red)')
    add_bullet_point(doc, 'What needs my attention? (Top 3-5 items requiring board action)')
    add_bullet_point(doc, 'Where are we trending? (Better, same, or worse than last period)')

    add_heading_with_style(doc, 'Principle 2: Exception-Based Reporting', level=3)
    doc.add_paragraph(
        'Don\'t show all 50 contracts—show the 5 that are off-track. Don\'t show all vendors—show the 3 with declining performance.'
    )

    add_heading_with_style(doc, 'Principle 3: Forward-Looking, Not Historical', level=3)
    doc.add_paragraph(
        'Executives care more about "What will happen" than "What happened." Show projected completion dates, '
        'forecast budget overruns, and predicted risks.'
    )

    add_heading_with_style(doc, 'Principle 4: Decision-Centric Design', level=3)
    doc.add_paragraph(
        'Every section should map to a decision. Examples:'
    )
    add_bullet_point(doc, 'Section: "Projects Requiring Funding Decision" → Decision: Approve or defer additional funding')
    add_bullet_point(doc, 'Section: "Vendors Requiring Review" → Decision: Renew, renegotiate, or replace')
    add_bullet_point(doc, 'Section: "Pipeline for Next Fiscal Year" → Decision: Which projects to prioritize in budget')

    add_heading_with_style(doc, 'Principle 5: Storytelling Through Data', level=3)
    doc.add_paragraph(
        'Use narrative annotations, auto-generated insights, and comparative context to transform raw numbers into actionable intelligence.'
    )

    doc.add_paragraph()

    add_heading_with_style(doc, 'Information Layering (Summary → Drill-Down)', level=2)

    doc.add_paragraph('Optimal executive dashboards use a three-tier information architecture:')

    para = doc.add_paragraph()
    run = para.add_run('Tier 1: Executive Summary Screen (30-Second View)')
    run.bold = True

    add_bullet_point(doc, 'Overall Portfolio Health: One score (0-100) with color indicator')
    add_bullet_point(doc, 'Top 3 Items Requiring Board Action: Red flags with decision prompt')
    add_bullet_point(doc, 'Trend Indicator: "Portfolio health improved 5 points this month"')
    add_bullet_point(doc, 'Key Metric Summary: 3-4 strategic KPIs (budget variance, on-time %, at-risk projects)')

    para = doc.add_paragraph()
    run = para.add_run('Tier 2: Strategic Overview (2-3 Minute View)')
    run.bold = True

    add_bullet_point(doc, 'Top 10 Projects by Strategic Importance: Ranked table with health status')
    add_bullet_point(doc, 'Spending Trends: 12-month chart with annotations explaining variances')
    add_bullet_point(doc, 'Vendor Performance Summary: Top/bottom performers with action recommendations')
    add_bullet_point(doc, 'Pipeline View: Upcoming projects by fiscal year with capacity analysis')

    para = doc.add_paragraph()
    run = para.add_run('Tier 3: Detailed Analytics (On-Demand Drill-Down)')
    run.bold = True

    add_bullet_point(doc, 'Click any summary metric to see underlying details')
    add_bullet_point(doc, 'Full contract list with advanced filtering')
    add_bullet_point(doc, 'Historical analysis, compliance details, audit trails')
    add_bullet_point(doc, 'Export to PDF/Excel for deep-dive analysis')

    doc.add_paragraph(
        'Navigation: Use tabbed interface or progressive disclosure. Default view is Tier 1. Executives click to expand details only when needed.'
    )

    doc.add_paragraph()

    add_heading_with_style(doc, 'Key Sections an Executive Dashboard Must Contain', level=2)

    add_heading_with_style(doc, '1. Hero Section: Overall Status', level=3)
    add_bullet_point(doc, 'Large, prominent overall health score')
    add_bullet_point(doc, 'Trend arrow (up/down/flat) with % change')
    add_bullet_point(doc, 'One-sentence auto-generated summary: "Portfolio is stable. 3 projects require attention."')

    add_heading_with_style(doc, '2. Action Required Section', level=3)
    add_bullet_point(doc, 'Top 3-5 items needing board decision, prioritized by urgency and impact')
    add_bullet_point(doc, 'Clear call-to-action for each item')
    add_bullet_point(doc, 'Example: "Approve $2.3M additional funding for Belleview Elementary—project at risk without action by March 15"')

    add_heading_with_style(doc, '3. Portfolio Performance Summary', level=3)
    add_bullet_point(doc, '4-6 strategic KPIs with comparison to targets')
    add_bullet_point(doc, 'Color-coded (green = on target, yellow = warning, red = critical)')
    add_bullet_point(doc, 'Include: Budget variance, schedule variance, vendor performance, compliance rate')

    add_heading_with_style(doc, '4. Top Projects Dashboard', level=3)
    add_bullet_point(doc, 'Table of 10 highest-value or highest-risk projects')
    add_bullet_point(doc, 'Columns: Project name, budget, % complete, health score, status, trend')
    add_bullet_point(doc, 'Click-to-expand for full project details')

    add_heading_with_style(doc, '5. Spending & Budget Trends', level=3)
    add_bullet_point(doc, 'Line chart showing 12-month spending trend')
    add_bullet_point(doc, 'Annotated with major events: "June spike due to summer construction season"')
    add_bullet_point(doc, 'Forecast line showing projected spend for next 6 months')

    add_heading_with_style(doc, '6. Vendor Insights', level=3)
    add_bullet_point(doc, 'Top performing vendors (last 3 years)')
    add_bullet_point(doc, 'Vendors with declining performance trends')
    add_bullet_point(doc, 'Diversity metrics: % minority-owned, local, small business')

    add_heading_with_style(doc, '7. Risk & Compliance Summary', level=3)
    add_bullet_point(doc, 'Count of at-risk projects with breakdown by risk type')
    add_bullet_point(doc, 'Compliance rate (insurance, bonds, board approvals)')
    add_bullet_point(doc, 'Upcoming expirations and renewals (30/60/90 day view)')

    add_heading_with_style(doc, '8. Comparative Context', level=3)
    add_bullet_point(doc, 'Marion County vs. peer counties: spending per capita, on-time %, cost variance')
    add_bullet_point(doc, 'Year-over-year comparison: "15% improvement in on-time delivery vs. last year"')

    doc.add_paragraph()

    add_heading_with_style(doc, 'How Executives Should Consume It in Under 5 Minutes', level=2)

    para = doc.add_paragraph()
    run = para.add_run('0:00 - 0:30 (30 seconds): Scan Hero Section')
    run.bold = True
    add_bullet_point(doc, 'Read overall health score and status')
    add_bullet_point(doc, 'Note trend direction (improving/declining)')
    add_bullet_point(doc, 'Read auto-generated one-sentence summary')

    para = doc.add_paragraph()
    run = para.add_run('0:30 - 1:30 (1 minute): Review Action Required Section')
    run.bold = True
    add_bullet_point(doc, 'Scan list of 3-5 items requiring board decision')
    add_bullet_point(doc, 'Identify which items are on today\'s agenda')
    add_bullet_point(doc, 'Note any items requiring follow-up with staff')

    para = doc.add_paragraph()
    run = para.add_run('1:30 - 3:00 (1.5 minutes): Review Portfolio Performance Summary')
    run.bold = True
    add_bullet_point(doc, 'Quick scan of 4-6 strategic KPIs')
    add_bullet_point(doc, 'Focus on red/yellow indicators')
    add_bullet_point(doc, 'Note any significant changes from last month')

    para = doc.add_paragraph()
    run = para.add_run('3:00 - 4:30 (1.5 minutes): Scan Top Projects Table')
    run.bold = True
    add_bullet_point(doc, 'Review top 10 projects for any new red flags')
    add_bullet_point(doc, 'Note projects showing improvement or decline')
    add_bullet_point(doc, 'Identify projects to discuss in board meeting')

    para = doc.add_paragraph()
    run = para.add_run('4:30 - 5:00 (30 seconds): Review Spending Trend Chart')
    run.bold = True
    add_bullet_point(doc, 'Observe overall spending trajectory')
    add_bullet_point(doc, 'Note any unusual spikes or dips')
    add_bullet_point(doc, 'Validate forecast aligns with budget expectations')

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Outcome: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    para.add_run(
        'After 5 minutes, the executive should be able to answer: (1) Is our portfolio healthy overall? '
        '(2) What decisions do I need to make today? (3) Are there any emerging issues I should monitor? '
        '(4) How are we performing vs. peers and historical trends?'
    )

    doc.add_page_break()

    # ========================================
    # 5. RECOMMENDED KPIs & METRICS
    # ========================================
    add_heading_with_style(doc, '5. Recommended KPIs & Metrics (High Level)', level=1)

    add_heading_with_style(doc, 'Strategic KPIs Relevant to School Projects and Capital Planning', level=2)

    doc.add_paragraph(
        'Executive dashboards should track no more than 8-12 KPIs. The following are recommended for '
        'school district capital project oversight:'
    )

    add_heading_with_style(doc, 'Financial Performance KPIs', level=3)

    para = doc.add_paragraph()
    run = para.add_run('1. Portfolio Budget Variance (%)')
    run.bold = True
    add_bullet_point(doc, 'Definition: (Actual Spend - Budgeted Spend) / Budgeted Spend', level=1)
    add_bullet_point(doc, 'Target: <5% over budget', level=1)
    add_bullet_point(doc, 'Why it matters: Core indicator of fiscal discipline and project planning accuracy', level=1)

    para = doc.add_paragraph()
    run = para.add_run('2. Cost Per Square Foot (Capital Projects)')
    run.bold = True
    add_bullet_point(doc, 'Definition: Total construction cost (excluding land) / Total square footage', level=1)
    add_bullet_point(doc, 'Target: Benchmark against peer districts and inflation-adjusted historical costs', level=1)
    add_bullet_point(doc, 'Why it matters: Enables apples-to-apples comparison across projects and identifies cost inefficiencies', level=1)

    para = doc.add_paragraph()
    run = para.add_run('3. Forecast Budget at Completion')
    run.bold = True
    add_bullet_point(doc, 'Definition: Projected final cost based on current burn rate and remaining work', level=1)
    add_bullet_point(doc, 'Target: ≤ Original approved budget', level=1)
    add_bullet_point(doc, 'Why it matters: Forward-looking metric that predicts final costs, enabling proactive intervention', level=1)

    add_heading_with_style(doc, 'Schedule Performance KPIs', level=3)

    para = doc.add_paragraph()
    run = para.add_run('4. On-Time Delivery Rate (%)')
    run.bold = True
    add_bullet_point(doc, 'Definition: % of projects completed by original target date', level=1)
    add_bullet_point(doc, 'Target: ≥85%', level=1)
    add_bullet_point(doc, 'Why it matters: Schedule delays often cascade into cost overruns and community dissatisfaction', level=1)

    para = doc.add_paragraph()
    run = para.add_run('5. Average Project Delay (Days)')
    run.bold = True
    add_bullet_point(doc, 'Definition: Average delay across all late projects', level=1)
    add_bullet_point(doc, 'Target: <30 days', level=1)
    add_bullet_point(doc, 'Why it matters: Identifies systemic project management issues vs. one-off delays', level=1)

    add_heading_with_style(doc, 'Vendor & Procurement KPIs', level=3)

    para = doc.add_paragraph()
    run = para.add_run('6. Vendor Performance Score (Average)')
    run.bold = True
    add_bullet_point(doc, 'Definition: Weighted average of quality, timeliness, communication, and value ratings', level=1)
    add_bullet_point(doc, 'Target: ≥80/100', level=1)
    add_bullet_point(doc, 'Why it matters: Poor vendor performance is the #1 predictor of project failure', level=1)

    para = doc.add_paragraph()
    run = para.add_run('7. Local Vendor Participation Rate (%)')
    run.bold = True
    add_bullet_point(doc, 'Definition: % of contract value awarded to local (in-county) vendors', level=1)
    add_bullet_point(doc, 'Target: ≥30% (adjust based on local policy)', level=1)
    add_bullet_point(doc, 'Why it matters: Economic development and community engagement priority', level=1)

    para = doc.add_paragraph()
    run = para.add_run('8. Vendor Concentration Risk (%)')
    run.bold = True
    add_bullet_point(doc, 'Definition: % of portfolio value with single largest vendor', level=1)
    add_bullet_point(doc, 'Target: <25% (no single vendor should dominate)', level=1)
    add_bullet_point(doc, 'Why it matters: Over-reliance on one vendor creates supply chain and negotiation risks', level=1)

    add_heading_with_style(doc, 'Risk & Compliance KPIs', level=3)

    para = doc.add_paragraph()
    run = para.add_run('9. At-Risk Project Percentage (%)')
    run.bold = True
    add_bullet_point(doc, 'Definition: % of projects with health score <50', level=1)
    add_bullet_point(doc, 'Target: <15%', level=1)
    add_bullet_point(doc, 'Why it matters: Leading indicator of portfolio instability', level=1)

    para = doc.add_paragraph()
    run = para.add_run('10. Compliance Rate (%)')
    run.bold = True
    add_bullet_point(doc, 'Definition: % of contracts with all required approvals, insurance, and bonds in place', level=1)
    add_bullet_point(doc, 'Target: 100%', level=1)
    add_bullet_point(doc, 'Why it matters: Non-compliance creates legal and financial liability', level=1)

    add_heading_with_style(doc, 'Capital Planning KPIs', level=3)

    para = doc.add_paragraph()
    run = para.add_run('11. Capital Spend Per Student (District-Level)')
    run.bold = True
    add_bullet_point(doc, 'Definition: Total capital investment / Student enrollment (by district)', level=1)
    add_bullet_point(doc, 'Target: Equitable distribution across all districts (variance <20%)', level=1)
    add_bullet_point(doc, 'Why it matters: Demonstrates equity and prevents political criticism of favoritism', level=1)

    para = doc.add_paragraph()
    run = para.add_run('12. Pipeline Capacity Utilization (%)')
    run.bold = True
    add_bullet_point(doc, 'Definition: (Active + Planned Projects Value) / Annual Capital Budget Capacity', level=1)
    add_bullet_point(doc, 'Target: 80-95% (not under- or over-committed)', level=1)
    add_bullet_point(doc, 'Why it matters: Identifies if district has capacity to take on new projects', level=1)

    doc.add_paragraph()

    add_heading_with_style(doc, 'What Should Be Elevated to the Top', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Primary "Hero Metrics" (Most Prominent on Dashboard):')
    run.bold = True

    add_bullet_point(doc, 'Overall Portfolio Health Score (0-100)')
    add_bullet_point(doc, 'Portfolio Budget Variance (%)')
    add_bullet_point(doc, 'On-Time Delivery Rate (%)')
    add_bullet_point(doc, 'At-Risk Project Count (absolute number)')

    doc.add_paragraph('These four metrics provide a complete snapshot of portfolio status.')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('Secondary KPIs (Visible but Less Prominent):')
    run.bold = True

    add_bullet_point(doc, 'Vendor Performance Score')
    add_bullet_point(doc, 'Local Vendor Participation Rate')
    add_bullet_point(doc, 'Compliance Rate')
    add_bullet_point(doc, 'Pipeline Capacity Utilization')

    doc.add_paragraph()

    add_heading_with_style(doc, 'What Should Be Removed, Deprioritized, or Moved to Supporting Views', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Remove from Executive Dashboard (Move to Detailed Analytics):')
    run.bold = True

    add_bullet_point(doc, 'Individual milestone tracking (too granular)')
    add_bullet_point(doc, 'Detailed payment history (operational, not strategic)')
    add_bullet_point(doc, 'Full contract terms and clauses (legal detail)')
    add_bullet_point(doc, 'Audit log activity (compliance team concern, not board-level)')
    add_bullet_point(doc, 'Benchmarking against Coupa 2025 KPIs (interesting but not directly actionable for board)')

    para = doc.add_paragraph()
    run = para.add_run('Deprioritize (Show Only on Drill-Down):')
    run.bold = True

    add_bullet_point(doc, 'Change order details (unless excessive—flag only if 3+ change orders)')
    add_bullet_point(doc, 'Insurance and bond verification status (unless non-compliant)')
    add_bullet_point(doc, 'Individual vendor ratings (show aggregate only; detail on click)')
    add_bullet_point(doc, 'Historical audit trails (archive; not needed for forward-looking decisions)')

    para = doc.add_paragraph()
    run = para.add_run('Consolidate:')
    run.bold = True

    add_bullet_point(doc, 'Reduce number of charts on executive dashboard from 6+ to 2-3 maximum')
    add_bullet_point(doc, 'Combine multiple financial metrics into single "Budget Health" scorecard')
    add_bullet_point(doc, 'Merge compliance indicators into single compliance rate percentage')

    doc.add_page_break()

    # ========================================
    # 6. SOLUTION ROADMAP
    # ========================================
    add_heading_with_style(doc, '6. Solution Roadmap', level=1)

    doc.add_paragraph(
        'The following phased roadmap outlines the transformation from current state to executive-grade dashboard. '
        'Each phase is designed to deliver immediate value while building toward the complete vision.'
    )

    doc.add_paragraph()

    # ===== PHASE 1: SHORT-TERM =====
    add_heading_with_style(doc, 'Phase 1: Short-Term Quick Wins (1-3 Weeks)', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Priority Level: CRITICAL')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('1.1 Simplify Executive Dashboard Layout')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Reduce KPI count from 12+ to 4-6 most strategic metrics', level=1)
    add_bullet_point(doc, 'Implement visual hierarchy: Largest = most important (hero metrics top-left)', level=1)
    add_bullet_point(doc, 'Add white space—reduce data density by 40%', level=1)
    add_bullet_point(doc, 'Create clear sections with headers: "Action Required," "Portfolio Status," "Trends"', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Board members can scan dashboard in 2-3 minutes vs. current 15-20 minutes')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('1.2 Add Auto-Generated Executive Summary')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Implement text summary generator that produces 2-3 sentence overview', level=1)
    add_bullet_point(doc, 'Example: "Portfolio health is stable at 72/100. 3 projects require board attention due to budget overruns. Overall spending is 4% under budget this quarter."', level=1)
    add_bullet_point(doc, 'Display prominently at top of dashboard', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Provides immediate context without requiring executives to interpret raw numbers')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('1.3 Create "Action Required" Priority Section')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Implement algorithm to identify top 3-5 items requiring board decision', level=1)
    add_bullet_point(doc, 'Criteria: Projects >$5M with health score <50, projects requiring additional funding, contract renewals >$1M', level=1)
    add_bullet_point(doc, 'Display as prominent cards with clear call-to-action', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Board members immediately see what decisions they need to make—core value proposition of dashboard')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('1.4 Replace Sample Data with Real Data (Critical)')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Identify data sources for real Marion County school district contracts', level=1)
    add_bullet_point(doc, 'Document challenges with obtaining real data (access restrictions, data quality issues)', level=1)
    add_bullet_point(doc, 'Create data import pipeline for ongoing updates', level=1)
    add_bullet_point(doc, 'Replace all sample data in database', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Dashboard becomes credible and usable for actual board decisions—without real data, dashboard is a demo tool only')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('1.5 Separate Land Costs from Construction Costs')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Add "land_cost" field to contracts database schema', level=1)
    add_bullet_point(doc, 'Create calculated field: construction_cost = total_cost - land_cost', level=1)
    add_bullet_point(doc, 'Update "Cost Per Square Foot" KPI to use construction_cost only', level=1)
    add_bullet_point(doc, 'Add toggle to charts: "Show with/without land costs"', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Enables accurate project cost comparisons—critical for benchmarking and decision-making')

    doc.add_paragraph()
    doc.add_paragraph()

    # ===== PHASE 2: MID-TERM =====
    add_heading_with_style(doc, 'Phase 2: Mid-Term Structural Improvements (4-8 Weeks)', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Priority Level: HIGH')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 140, 0)

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('2.1 Implement Predictive Analytics')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Build "Forecast Budget at Completion" model based on burn rate analysis', level=1)
    add_bullet_point(doc, 'Add "Expected Completion Date" prediction using current progress velocity', level=1)
    add_bullet_point(doc, 'Create risk probability scoring: likelihood of project failure based on historical patterns', level=1)
    add_bullet_point(doc, 'Display predictions prominently with confidence intervals', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Transforms dashboard from backward-looking to forward-looking—enables proactive intervention before projects fail')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('2.2 Build Historical Vendor Analysis (20+ Years)')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Import historical contract and vendor data (20+ years if available)', level=1)
    add_bullet_point(doc, 'Calculate long-term vendor performance metrics: average cost variance, on-time %, reliability', level=1)
    add_bullet_point(doc, 'Create vendor track record dashboard showing historical performance', level=1)
    add_bullet_point(doc, 'Add "Vendor Recommendation" feature: highlight vendors with best long-term performance for similar projects', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Supports strategic vendor selection for major multi-year contracts—reduces risk of poor vendor performance')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('2.3 Create Capital Project Pipeline View')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Build pipeline dashboard showing planned projects by fiscal year', level=1)
    add_bullet_point(doc, 'Display: Project name, estimated cost, proposed start date, funding status', level=1)
    add_bullet_point(doc, 'Add capacity analysis: "You have $15M budget capacity remaining for FY26"', level=1)
    add_bullet_point(doc, 'Enable drag-and-drop prioritization of pipeline projects', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Enables multi-year capital planning and answers "Can we afford this project next year?"')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('2.4 Add Dollar Normalization (Inflation Adjustment)')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Integrate construction cost inflation index (e.g., ENR Construction Cost Index)', level=1)
    add_bullet_point(doc, 'Add "inflation_adjusted_cost" calculated field to all historical contracts', level=1)
    add_bullet_point(doc, 'Display historical cost comparisons in "constant dollars" (e.g., 2026 dollars)', level=1)
    add_bullet_point(doc, 'Add toggle: "Show actual costs / Show inflation-adjusted costs"', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Enables accurate historical cost comparisons—answers "Are we paying more or less than we did 10 years ago?"')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('2.5 Implement District-Level Equity Analysis')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Calculate capital spend per student by school district (elementary, middle, high school zones)', level=1)
    add_bullet_point(doc, 'Create geographic heat map showing investment distribution', level=1)
    add_bullet_point(doc, 'Add equity dashboard: "District A has received 35% more per-student investment than District B over 5 years"', level=1)
    add_bullet_point(doc, 'Provide equity score: How evenly are capital investments distributed?', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Addresses political/constituent concerns about fair resource distribution—critical for public trust')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('2.6 Build Project Risk Assessment Model')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Develop risk scoring model: Likelihood of failure × Financial impact', level=1)
    add_bullet_point(doc, 'Risk factors: Vendor performance, project complexity, budget size, timeline, change order history', level=1)
    add_bullet_point(doc, 'Display: "High Risk: 35% probability of >20% cost overrun"', level=1)
    add_bullet_point(doc, 'Create "Project Risk Matrix" quadrant chart: Impact vs. Likelihood', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Provides data-driven risk assessment for go/no-go funding decisions')

    doc.add_paragraph()
    doc.add_paragraph()

    # ===== PHASE 3: LONG-TERM =====
    add_heading_with_style(doc, 'Phase 3: Long-Term Advanced Analytics & Governance (9-16 Weeks)', level=2)

    para = doc.add_paragraph()
    run = para.add_run('Priority Level: MEDIUM (Future Enhancement)')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('3.1 Implement Scenario Planning & What-If Analysis')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Build scenario modeling tool: "What if we delay this project 6 months?"', level=1)
    add_bullet_point(doc, 'Cost impact modeling: "What if material costs increase 15%?"', level=1)
    add_bullet_point(doc, 'Portfolio optimization: "How should we prioritize projects to maximize ROI?"', level=1)
    add_bullet_point(doc, 'Budget constraint modeling: "If budget is cut 10%, which projects should we defer?"', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Supports strategic planning discussions—executives can model different scenarios before making decisions')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('3.2 Add AI-Powered Insights & Recommendations')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Integrate Claude API (Anthropic) for automated insight generation', level=1)
    add_bullet_point(doc, 'AI analyzes trends and generates natural language summaries', level=1)
    add_bullet_point(doc, 'Example: "Vendor X\'s performance has declined 15% over last 3 quarters. Consider competitive rebid for next contract."', level=1)
    add_bullet_point(doc, 'Provide AI-generated recommendations: "Based on historical data, approve Project Y but defer Project Z until Q3"', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Transforms dashboard into intelligent advisor—proactively surfaces insights executives might miss')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('3.3 Build Automated Board Report Generator')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Create "Generate Board Report" button that produces formatted PDF', level=1)
    add_bullet_point(doc, 'Report includes: Executive summary, top projects, risks, recommendations, supporting charts', level=1)
    add_bullet_point(doc, 'Customizable: Select which projects/topics to include in report', level=1)
    add_bullet_point(doc, 'Branded template matching board presentation standards', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Eliminates manual report preparation—staff can generate board-ready reports in 5 minutes vs. hours')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('3.4 Implement Mobile-Optimized Executive View')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Create mobile app or mobile-optimized web view', level=1)
    add_bullet_point(doc, 'Simplified layout for smartphone screens', level=1)
    add_bullet_point(doc, 'Push notifications for critical alerts: "Project X exceeded budget by $500K"', level=1)
    add_bullet_point(doc, 'Offline mode: Download executive summary for review without internet', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Executives can review portfolio status anytime, anywhere—increases engagement and responsiveness')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('3.5 Establish Governance & Continuous Improvement Framework')
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run('Actions:')
    run.bold = True
    add_bullet_point(doc, 'Create dashboard governance committee (board members + staff)', level=1)
    add_bullet_point(doc, 'Quarterly review: "Is the dashboard meeting board needs?"', level=1)
    add_bullet_point(doc, 'Implement feedback mechanism: Executives can suggest improvements directly in dashboard', level=1)
    add_bullet_point(doc, 'Track dashboard engagement analytics: Which sections are most viewed? Which are ignored?', level=1)
    add_bullet_point(doc, 'Continuous refinement based on actual usage patterns', level=1)

    para = doc.add_paragraph()
    run = para.add_run('Expected Executive Value:')
    run.bold = True
    para.add_run(' Ensures dashboard remains relevant and valuable over time—avoids becoming stale or obsolete')

    doc.add_page_break()

    # ========================================
    # CONCLUSION
    # ========================================
    add_heading_with_style(doc, 'Conclusion', level=1)

    doc.add_paragraph(
        'The Marion County Contract Oversight System possesses a strong technical foundation and comprehensive '
        'data infrastructure. The transformation to an executive-grade Board Decision Dashboard is not a rebuild—it '
        'is a focused refinement of the presentation layer and information architecture.'
    )

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('Key Success Factors:')
    run.bold = True

    add_bullet_point(doc, 'Prioritize clarity over comprehensiveness—executives need insights, not data dumps')
    add_bullet_point(doc, 'Implement quick wins first (Phase 1) to demonstrate immediate value')
    add_bullet_point(doc, 'Replace sample data with real data before any board presentation—credibility is paramount')
    add_bullet_point(doc, 'Separate land costs from construction costs—foundational for accurate analysis')
    add_bullet_point(doc, 'Add forward-looking metrics—boards care more about what will happen than what happened')
    add_bullet_point(doc, 'Establish governance framework to ensure continuous improvement')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('Recommended Next Steps:')
    run.bold = True

    add_bullet_point(doc, 'Review this gap analysis with project stakeholders and board leadership')
    add_bullet_point(doc, 'Prioritize Phase 1 quick wins for immediate implementation')
    add_bullet_point(doc, 'Identify data sources for replacing sample data with real Marion County data')
    add_bullet_point(doc, 'Schedule demo session with school board executives to gather direct feedback')
    add_bullet_point(doc, 'Allocate resources for phased roadmap execution')

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('Final Observation:')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph(
        'This dashboard has the potential to set a new standard for government transparency and data-driven '
        'decision-making in public education infrastructure. With focused refinement guided by this gap analysis, '
        'it can become the authoritative tool for all capital project decisions, demonstrating best practices '
        'in fiscal accountability and strategic planning.'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run('— End of Gap Analysis Report —')
    footer_run.font.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

if __name__ == "__main__":
    print("Generating Executive Dashboard Gap Analysis document...")
    doc = create_gap_analysis_document()

    output_path = "Executive_Dashboard_Gap_Analysis.docx"
    doc.save(output_path)

    print(f"[SUCCESS] Document successfully generated: {output_path}")
    print(f"  File size: {round(os.path.getsize(output_path) / 1024, 1)} KB")
    print("\nDocument contains:")
    print("  1. Executive Summary")
    print("  2. Current State Assessment")
    print("  3. Gap Analysis (5 dimensions)")
    print("  4. Ideal Executive Dashboard Vision")
    print("  5. Recommended KPIs & Metrics")
    print("  6. Solution Roadmap (3 phases)")
    print("\nReady for review and presentation to school board leadership.")
